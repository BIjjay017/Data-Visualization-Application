import pandas as pd
import io
from PyPDF2 import PdfReader
from docx import Document
import openpyxl

def parse_csv(file_content):
    """Parse CSV file"""
    return pd.read_csv(io.BytesIO(file_content))

def parse_excel(file_content):
    """Parse Excel files (XLS, XLSX)"""
    return pd.read_excel(io.BytesIO(file_content))

def parse_pdf(file_content):
    """Parse PDF file and extract text as structured data"""
    pdf_reader = PdfReader(io.BytesIO(file_content))
    
    # Extract text from all pages
    text_data = []
    for page in pdf_reader.pages:
        text = page.extract_text()
        if text.strip():
            # Try to parse as table-like structure
            lines = text.strip().split('\n')
            for line in lines:
                # Split by multiple spaces or tabs
                row = [cell.strip() for cell in line.split('  ') if cell.strip()]
                if row:
                    text_data.append(row)
    
    if not text_data:
        raise ValueError("No extractable data found in PDF")
    
    # Try to create DataFrame
    # Assume first row is header
    if len(text_data) > 1:
        df = pd.DataFrame(text_data[1:], columns=text_data[0])
    else:
        df = pd.DataFrame(text_data)
    
    return df

def parse_word(file_content):
    """Parse Word document and extract tables"""
    doc = Document(io.BytesIO(file_content))
    
    # Extract tables from document
    all_tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        
        if table_data:
            # Assume first row is header
            if len(table_data) > 1:
                df = pd.DataFrame(table_data[1:], columns=table_data[0])
            else:
                df = pd.DataFrame(table_data)
            all_tables.append(df)
    
    if not all_tables:
        # If no tables, try to extract text as simple data
        text_data = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_data.append([text])
        
        if text_data:
            df = pd.DataFrame(text_data, columns=['Content'])
            return df
        else:
            raise ValueError("No extractable data found in Word document")
    
    # Return the first table (or combine multiple tables if needed)
    return all_tables[0]

def parse_file(file_content, filename):
    """
    Parse file based on extension and return pandas DataFrame
    """
    extension = filename.lower().split('.')[-1]
    
    try:
        if extension == 'csv':
            return parse_csv(file_content)
        elif extension in ['xls', 'xlsx']:
            return parse_excel(file_content)
        elif extension == 'pdf':
            return parse_pdf(file_content)
        elif extension in ['doc', 'docx']:
            return parse_word(file_content)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    except Exception as e:
        raise ValueError(f"Error parsing {extension.upper()} file: {str(e)}")
